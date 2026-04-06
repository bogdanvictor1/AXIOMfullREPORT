import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from PIL import Image

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -------------------------
# Configuração global (definida pela GUI antes de cada execução)
# -------------------------
MAX_PAGES = 0  # 0 = todas as páginas

# -------------------------
# Helpers
# -------------------------
def safe_text(elem):
    return elem.text if elem is not None and elem.text else ''

def find_fragment_text(hit, *names):
    for frag in hit.findall('Fragment'):
        if frag.get('name', '').strip() in names:
            return frag.text or ''
    return ''

def find_fragment_path(hit, names):
    # Legado: texto que contenha 'Attachments' (ou attachment="True")
    for frag in hit.findall('Fragment'):
        if frag.get('name') in names and frag.text and "Attachments" in frag.text:
            return frag.text
    for frag in hit.findall('Fragment'):
        if frag.get('attachment', '').lower() == 'true' and frag.text and "Attachments" in frag.text:
            return frag.text
    for frag in hit.findall('Fragment'):
        if frag.text and "Attachments" in frag.text:
            return frag.text
    return None

def get_field_value(hit, aliases):
    if 'sequenceNumber' in aliases:
        return hit.get('sequenceNumber', '')
    return find_fragment_text(hit, *aliases)

def _normalize_relpath(text):
    return text.replace('\\\\', os.sep).replace('/', os.sep).lstrip('.' + os.sep)

def safe_slug(text):
    t = text.lower()
    trans = str.maketrans("áàâãäéèêëíìîïóòôõöúùûüçñ", "aaaaaeeeeiiiiooooouuuucn")
    t = t.translate(trans)
    out = []
    for ch in t:
        out.append(ch if ch.isalnum() else '_')
    s = ''.join(out)
    while '__' in s:
        s = s.replace('__', '_')
    return s.strip('_')

# -------------------------
# Métricas e limites
# -------------------------
def _image_metrics_mm(img_path, fallback_dpi=96):
    """
    Retorna (w_mm, h_mm, w_px, h_px) usando DPI do arquivo se houver; fallback 96 dpi.
    """
    with Image.open(img_path) as im:
        w_px, h_px = im.size
        dpi = im.info.get('dpi', (fallback_dpi, fallback_dpi))[0] or fallback_dpi
    w_mm = w_px / float(dpi) * 25.4
    h_mm = h_px / float(dpi) * 25.4
    return w_mm, h_mm, w_px, h_px

def _fit_box_if_needed(img_path, max_w_mm, max_h_mm):
    """
    Decide pelo tamanho de exibição:
      - None => cabe nos limites (usar tamanho original em mm)
      - {'width_mm': X} / {'height_mm': Y} => reduzir proporcionalmente para caber
    """
    w_mm, h_mm, w_px, h_px = _image_metrics_mm(img_path)
    if w_mm <= max_w_mm and h_mm <= max_h_mm:
        return None, w_mm, h_mm
    r = w_px / float(h_px)
    box_r = max_w_mm / float(max_h_mm)
    if r >= box_r:
        return {'width_mm': max_w_mm}, w_mm, h_mm
    else:
        return {'height_mm': max_h_mm}, w_mm, h_mm

# -------------------------
# Borda NA IMAGEM (stroke) — altera o XML do desenho
# -------------------------
def _add_image_stroke_to_run(run, color="000000", width_pt=0.75):
    """
    Aplica uma borda (linha) ao redor da imagem no último desenho do run:
    - Encontra w:drawing -> pic:pic -> pic:spPr
    - Cria a:ln com a:solidFill/a:srgbClr e largura em EMUs (1 pt = 12700 EMUs)
    """
    EMUS_PER_PT = 12700
    width_emus = int(round(width_pt * EMUS_PER_PT))

    # Namespaces usados no DrawingML
    ns = {
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    }

    # Localiza o último <pic:pic> dentro do run
    try:
        pics = run._r.xpath('.//pic:pic', namespaces=ns)
        if not pics:
            return
        pic = pics[-1]
        spPr = pic.find('pic:spPr', namespaces=ns)
        if spPr is None:
            # cria spPr se não existir
            spPr = OxmlElement(qn('pic:spPr'))
            pic.append(spPr)

        # a:ln (linha/contorno)
        ln = spPr.find(qn('a:ln'))
        if ln is None:
            ln = OxmlElement(qn('a:ln'))
            spPr.append(ln)
        ln.set('w', str(width_emus))

        # a:solidFill/a:srgbClr
        solid = ln.find(qn('a:solidFill'))
        if solid is None:
            solid = OxmlElement(qn('a:solidFill'))
            ln.append(solid)
        srgb = solid.find(qn('a:srgbClr'))
        if srgb is None:
            srgb = OxmlElement(qn('a:srgbClr'))
            solid.append(srgb)
        srgb.set(qn('a:val'), color)

        # (opcional) ponta/composto
        # ln.set(qn('a:cmpd'), 'sng')  # linha simples
        # ln.set(qn('a:cap'), 'flat')
    except Exception:
        # Se algo falhar silenciosamente, mantemos a imagem sem borda para não quebrar o fluxo
        pass

# -------------------------
# Inserção de imagens FORA de tabela/célula
# -------------------------
def insert_doc_image(doc, img_path, max_w_mm=155, max_h_mm=215):
    """
    Anexos de documentos (Word/PDF/Excel):
    - Fora de tabela/célula.
    - Qualidade/tamanho original; só reduz para caber em 155 x 215 mm.
    - Borda 0,75 pt na própria imagem. Sem legenda.
    """
    reduce_spec, w_orig_mm, h_orig_mm = _fit_box_if_needed(img_path, max_w_mm, max_h_mm)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run()
    if reduce_spec is None:
        run.add_picture(img_path, width=Mm(w_orig_mm))  # mantém o tamanho "original" de exibição
    else:
        if 'width_mm' in reduce_spec and reduce_spec['width_mm'] is not None:
            run.add_picture(img_path, width=Mm(reduce_spec['width_mm']))
        else:
            run.add_picture(img_path, height=Mm(reduce_spec['height_mm']))

    _add_image_stroke_to_run(run, color="000000", width_pt=0.75)

def insert_legacy_image(doc, img_path, max_w_mm=155, max_h_mm=215):
    """
    Artefatos legacy (Imagens/Thumbcache/Vídeos):
    - Fora de tabela/célula.
    - Qualidade/tamanho original; só reduz para caber em 155 x 215 mm.
    - Borda 0,75 pt na imagem. Com legenda (Calibri 10) em parágrafo separado.
    """
    reduce_spec, w_orig_mm, h_orig_mm = _fit_box_if_needed(img_path, max_w_mm, max_h_mm)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run()
    if reduce_spec is None:
        run.add_picture(img_path, width=Mm(w_orig_mm))
    else:
        if 'width_mm' in reduce_spec and reduce_spec['width_mm'] is not None:
            run.add_picture(img_path, width=Mm(reduce_spec['width_mm']))
        else:
            run.add_picture(img_path, height=Mm(reduce_spec['height_mm']))

    _add_image_stroke_to_run(run, color="000000", width_pt=0.75)

    # Legenda (sem borda)
    cap = doc.add_paragraph(os.path.basename(img_path))
    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    if cap.runs:
        cap.runs[0].font.name = 'Calibri'
        cap.runs[0].font.size = Pt(10)


# -------------------------
# Funções para geração automática de imagens
# -------------------------
import fitz  # PyMuPDF — pip install pymupdf (sem dependencias externas)

def generate_images_from_pdf(pdf_path, output_dir, dpi=200, max_pages=0):
    """
    Converte paginas do PDF em JPEG usando PyMuPDF (sem Poppler ou qualquer
    dependencia externa — apenas: pip install pymupdf).
    max_pages=0 => todas as paginas; max_pages>0 => somente as primeiras N paginas.
    Retorna lista de caminhos das imagens geradas.
    """
    if not os.path.isfile(pdf_path):
        return []
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    zoom = dpi / 72.0  # PDF usa 72 dpi como base
    mat = fitz.Matrix(zoom, zoom)

    img_paths = []
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]

    with fitz.open(pdf_path) as pdf:
        total = len(pdf)
        limit = total if (not max_pages or max_pages <= 0) else min(max_pages, total)
        for i in range(limit):
            page = pdf[i]
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_filename = f"{base_name}_pg{i + 1}.jpg"
            img_path = os.path.join(output_dir, img_filename)
            pix.save(img_path)
            img_paths.append(img_path)

    return img_paths

def _docx_to_pdf_via_com(docx_path, pdf_path):
    """
    Converte .doc/.docx para PDF usando Microsoft Word via win32com (Windows).
    Requer: pip install pywin32   +   Microsoft Word instalado.
    """
    import win32com.client  # type: ignore

    docx_abs = os.path.abspath(docx_path)
    pdf_abs  = os.path.abspath(pdf_path)

    word = None
    doc  = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_abs, ReadOnly=True)
        # wdFormatPDF = 17
        doc.SaveAs2(pdf_abs, FileFormat=17)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def generate_images_from_word(docx_path, output_dir, dpi=200, max_pages=0):
    """
    Estratégia para Word (.doc/.docx) no Windows:
      1. Busca <nome_sem_ext>.pdf no mesmo diretório do .docx
      2. Se não achar, converte via Microsoft Word (win32com)
      3. Converte o PDF resultante em imagens JPEG

    Requisitos para a conversão automática:
      - Microsoft Word instalado
      - pip install pywin32
    """
    if not os.path.isfile(docx_path):
        return [], "Arquivo não encontrado: " + docx_path

    file_dir = os.path.dirname(docx_path)
    stem     = os.path.splitext(os.path.basename(docx_path))[0]
    candidate_pdf = os.path.join(file_dir, stem + ".pdf")

    if os.path.isfile(candidate_pdf):
        pdf_to_use = candidate_pdf
        msg = None
    else:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        converted_pdf = os.path.join(output_dir, stem + "_converted.pdf")
        try:
            _docx_to_pdf_via_com(docx_path, converted_pdf)
            pdf_to_use = converted_pdf
            msg = f"PDF não encontrado para '{os.path.basename(docx_path)}'; convertido via Word."
        except ImportError:
            return [], (
                f"Não foi possível converter '{os.path.basename(docx_path)}': módulo 'win32com' não encontrado. "
                f"Execute: pip install pywin32"
            )
        except Exception as e:
            return [], f"Falha ao converter '{os.path.basename(docx_path)}' via Word COM: {e}"

    imgs = generate_images_from_pdf(pdf_to_use, output_dir, dpi=dpi, max_pages=max_pages)
    return imgs, msg


def generate_images_from_excel(xlsx_path, output_dir, dpi=200, max_pages=0, log_callback=None):
    """
    Estratégia para Excel (.xls/.xlsx):
      1. Busca <nome_com_ext>.pdf  ex.: arquivo.xlsx.pdf
      2. Busca <nome_sem_ext>.pdf  ex.: arquivo.pdf  (fallback)
      3. Se não achar nenhum, loga aviso e retorna lista vazia
    """
    if not os.path.isfile(xlsx_path):
        return []

    file_dir = os.path.dirname(xlsx_path)
    basename = os.path.basename(xlsx_path)          # arquivo.xlsx
    stem = os.path.splitext(basename)[0]            # arquivo

    candidates = [
        os.path.join(file_dir, basename + ".pdf"),  # arquivo.xlsx.pdf
        os.path.join(file_dir, stem + ".pdf"),       # arquivo.pdf
    ]

    pdf_to_use = next((p for p in candidates if os.path.isfile(p)), None)

    if pdf_to_use is None:
        if log_callback:
            log_callback(f"Aviso: nenhum PDF correspondente encontrado para '{basename}'. Pulando imagens.")
        return []

    return generate_images_from_pdf(pdf_to_use, output_dir, dpi=dpi, max_pages=max_pages)


def _pptx_to_pdf_via_com(pptx_path, pdf_path):
    """
    Converte .ppt/.pptx para PDF usando Microsoft PowerPoint via win32com (Windows).
    Requer: pip install pywin32   +   Microsoft PowerPoint instalado.
    Lança exceção em caso de falha.
    """
    import win32com.client  # type: ignore

    pptx_abs = os.path.abspath(pptx_path)
    pdf_abs  = os.path.abspath(pdf_path)

    powerpoint = None
    prs        = None
    try:
        powerpoint = win32com.client.Dispatch("PowerPoint.Application")
        powerpoint.Visible = 1   # PowerPoint exige janela visível para exportar PDF
        prs = powerpoint.Presentations.Open(pptx_abs, ReadOnly=True, Untitled=False, WithWindow=False)
        # ppSaveAsPDF = 32
        prs.SaveAs(pdf_abs, 32)
    finally:
        if prs is not None:
            try:
                prs.Close()
            except Exception:
                pass
        if powerpoint is not None:
            try:
                powerpoint.Quit()
            except Exception:
                pass


def generate_images_from_powerpoint(pptx_path, output_dir, dpi=200, max_pages=0, log_callback=None):
    """
    Estratégia para PowerPoint (.ppt/.pptx) no Windows:
      1. Busca <nome_sem_ext>.pdf no mesmo diretório do arquivo
      2. Busca <nome_com_ext>.pdf  ex.: arquivo.pptx.pdf
      3. Se não achar, converte via Microsoft PowerPoint (win32com)
      4. Converte o PDF resultante em imagens JPEG respeitando max_pages

    Requisitos para a conversão automática:
      - Microsoft PowerPoint instalado
      - pip install pywin32
    """
    if not os.path.isfile(pptx_path):
        return [], f"Arquivo não encontrado: {pptx_path}"

    file_dir = os.path.dirname(pptx_path)
    basename = os.path.basename(pptx_path)
    stem     = os.path.splitext(basename)[0]

    # 1) Busca PDF pré-existente
    candidates = [
        os.path.join(file_dir, stem + ".pdf"),      # arquivo.pdf
        os.path.join(file_dir, basename + ".pdf"),  # arquivo.pptx.pdf
    ]
    pdf_to_use = next((p for p in candidates if os.path.isfile(p)), None)

    if pdf_to_use:
        msg = None
    else:
        # 2) Converte via PowerPoint COM
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        converted_pdf = os.path.join(output_dir, stem + "_converted.pdf")
        try:
            _pptx_to_pdf_via_com(pptx_path, converted_pdf)
            pdf_to_use = converted_pdf
            msg = f"PDF não encontrado para '{basename}'; convertido via PowerPoint."
        except ImportError:
            return [], (
                f"Não foi possível converter '{basename}': módulo 'win32com' não encontrado. "
                f"Execute: pip install pywin32"
            )
        except Exception as e:
            return [], f"Falha ao converter '{basename}' via PowerPoint COM: {e}"

    imgs = generate_images_from_pdf(pdf_to_use, output_dir, dpi=dpi, max_pages=max_pages)
    return imgs, msg


def insert_csv_content(doc, csv_path, max_rows=200):
    """
    Lê um arquivo CSV e insere seu conteúdo como tabela no documento.
    Limita a max_rows linhas para não gerar documentos inviáveis.
    """
    import csv as csv_mod

    if not os.path.isfile(csv_path):
        p = doc.add_paragraph(f"[Arquivo CSV não encontrado: {os.path.basename(csv_path)}]")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        return

    # Detecta encoding
    encodings = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
    rows = []
    used_enc = 'utf-8'
    for enc in encodings:
        try:
            with open(csv_path, newline='', encoding=enc) as f:
                reader = csv_mod.reader(f)
                rows = [r for r in reader]
            used_enc = enc
            break
        except Exception:
            rows = []

    if not rows:
        p = doc.add_paragraph(f"[Não foi possível ler o CSV: {os.path.basename(csv_path)}]")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        return

    truncated = len(rows) > max_rows
    rows = rows[:max_rows]
    num_cols = max(len(r) for r in rows) if rows else 1

    tbl = doc.add_table(rows=0, cols=num_cols)
    tbl.style = 'Table Grid'
    tbl.autofit = True

    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.add_row().cells
        for c_idx in range(num_cols):
            val = row_data[c_idx] if c_idx < len(row_data) else ''
            cell = row_cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True

    if truncated:
        note = doc.add_paragraph(
            f"[Exibindo as primeiras {max_rows} linhas de {os.path.basename(csv_path)}]"
        )
        note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        note.paragraph_format.space_before = Pt(2)
        note.paragraph_format.space_after = Pt(6)
        if note.runs:
            note.runs[0].font.name = 'Calibri'
            note.runs[0].font.size = Pt(9)
            note.runs[0].italic = True


def resolve_pdf_by_hyperlink(hit, xml_dir, log_callback=None):
    """
    Para artefatos PDF: localiza o arquivo .pdf a partir do caminho local
    armazenado nos Fragments do hit (hiperlink / campo 'Arquivo' / 'File').
    Retorna o caminho absoluto ou None.
    """
    for frag in hit.findall('Fragment'):
        txt = frag.text or ''
        if not txt.strip():
            continue
        # Normaliza separadores e remove prefixo relativo
        norm = _normalize_relpath(txt.strip())
        abs_path = os.path.join(xml_dir, norm)
        if os.path.isfile(abs_path) and abs_path.lower().endswith('.pdf'):
            return abs_path
        # Às vezes o fragmento traz só o nome do arquivo, sem caminho
        if os.path.isfile(os.path.join(xml_dir, os.path.basename(norm))):
            candidate = os.path.join(xml_dir, os.path.basename(norm))
            if candidate.lower().endswith('.pdf'):
                return candidate
    return None
# -------------------------
# Multi-sequencial (.jpg/.jpeg)
# -------------------------
def list_sequential_jpegs(base_abs_path):
    """
    Retorna lista ordenada de caminhos .jpg/.jpeg:
      - <base>_<n>.jpg
      - <base>.<ext>_<n>.jpg  (ex.: base.pdf_01.jpg / base.docx_1.jpg / base.xlsx_02.jpg)
    Quando o fragmento não traz extensão, tentamos .pdf .doc .docx .xls .xlsx.
    """
    base_dir = os.path.dirname(base_abs_path)
    base_name = os.path.basename(base_abs_path)
    stem, ext = os.path.splitext(base_name)

    stems = set()
    if base_name:
        stems.add(base_name)
    if stem:
        stems.add(stem)
    if not ext:
        for e in ('.pdf', '.doc', '.docx', '.xls', '.xlsx'):
            stems.add(stem + e)
    stems = list(stems)

    results = []

    def scan_dir(directory):
        try:
            for fn in os.listdir(directory):
                for s in stems:
                    if re.match(rf'^{re.escape(s)}_(\d+)\.(?:jpe?g)$', fn, flags=re.IGNORECASE):
                        idx = int(re.findall(r'_(\d+)\.(?:jpe?g)$', fn, flags=re.IGNORECASE)[0])
                        results.append((idx, os.path.join(directory, fn)))
        except Exception:
            pass

    if os.path.isdir(base_dir):
        scan_dir(base_dir)

    if not results:
        root_dir = os.path.dirname(base_dir)
        attach_dir = os.path.join(root_dir, 'Attachments')
        if os.path.isdir(attach_dir):
            for root, _, _files in os.walk(attach_dir):
                depth = os.path.relpath(root, attach_dir).count(os.sep)
                if depth > 2:
                    continue
                scan_dir(root)

    results.sort(key=lambda x: x[0])
    return [p for _, p in results]

# -------------------------
# Tabela de metadados
# -------------------------
def set_col_widths(table, widths_mm):
    for col_idx, width in enumerate(widths_mm):
        for row in table.rows:
            cell = row.cells[col_idx]
            cell.width = Mm(width)
        if col_idx < len(table.columns):
            table.columns[col_idx].width = Mm(width)

def style_table(table):
    table.style = 'Table Grid'
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if not cell.paragraphs:
                cell.add_paragraph()
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if not p.runs:
                    p.add_run()
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0,0,0)
                    if idx == 0:
                        run.bold = True

# -------------------------
# Mapeamento
# -------------------------
def build_mapping():
    return {
        'Pictures': {
            'aliases': ['Pictures', 'Imagens'],
            'title': 'Imagens',
            'artifact_type': 'legacy',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Nome do arquivo', ('File Name', 'Nome do arquivo')),
                ('Data/hora da criação - UTC+00:00 (dd/MM/yyyy)', ('Created Date/Time - UTC+00:00 (dd/MM/yyyy)', 'Data/hora da criação - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora da última modificação - UTC+00:00 (dd/MM/yyyy)', ('Last Modified Date/Time - UTC+00:00 (dd/MM/yyyy)', 'Data/hora da última modificação - UTC+00:00 (dd/MM/yyyy)')),
                ('Tamanho (Bytes)', ('Size (Bytes)', 'Tamanho (Bytes)')),
                ('Dados Exif', ('Exif Data', 'Dados Exif')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': ['Image', 'Imagem'],
            'multi_seq_jpg': False
        },
        'Thumbcache Pictures': {
            'aliases': ['Thumbcache Pictures', 'Imagens do Thumbcache'],
            'title': 'Imagens do Thumbcache',
            'artifact_type': 'legacy',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Arquivo de miniatura', ('Thumbnail File', 'Arquivo de miniatura')),
                ('Tamanho (Bytes)', ('Size (Bytes)', 'Tamanho (Bytes)')),
                ('Nome do arquivo', ('File Name', 'Nome do arquivo')),
                ('Caminho do arquivo', ('File Path', 'Caminho do arquivo')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': ['Picture', 'Imagem'],
            'multi_seq_jpg': False
        },
        'Videos': {
            'aliases': ['Videos', 'Vídeos'],
            'title': 'Vídeos',
            'artifact_type': 'legacy',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Nome do arquivo', ('File Name', 'Nome do arquivo')),
                ('Data/hora da criação - UTC+00:00 (dd/MM/yyyy)', ('Created Date/Time - UTC+00:00 (dd/MM/yyyy)', 'Data/hora da criação - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora da última modificação - UTC+00:00 (dd/MM/yyyy)', ('Last Modified Date/Time - UTC+00:00 (dd/MM/yyyy)', 'Data/hora da última modificação - UTC+00:00 (dd/MM/yyyy)')),
                ('Tamanho do arquivo (Bytes)', ('File Size (Bytes)', 'Tamanho (Bytes)')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Dados Exif', ('Exif Data', 'Dados Exif')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': ['Image', 'Imagem'],
            'multi_seq_jpg': False
        },
        'Microsoft Word Documents': {
            'aliases': ['Documentos do Microsoft Word', 'Microsoft Word Documents'],
            'title': 'Documentos do Microsoft Word',
            'artifact_type': 'word',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Nome do arquivo', ('Filename', 'Nome do arquivo')),
                ('Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Last Modified Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Created Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Tamanho (Bytes)', ('Size (Bytes)', 'File Size (Bytes)', 'Tamanho (Bytes)')),
                ('Autores', ('Authors', 'Autores')),
                ('Último autor', ('Last Author', 'Último autor')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': ['Arquivo', 'File'],
            'multi_seq_jpg': True
        },
        'PDF Documents': {
            'aliases': ['Documentos PDF', 'PDF Documents'],
            'title': 'Documentos PDF',
            'artifact_type': 'pdf',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Nome do arquivo', ('Filename', 'Nome do arquivo')),
                ('Título', ('Title', 'Título')),
                ('Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Last Modified Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Created Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Tamanho (Bytes)', ('Size (Bytes)', 'File Size (Bytes)', 'Tamanho (Bytes)')),
                ('Autores', ('Authors', 'Autores')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': ['Arquivo', 'File'],
            'multi_seq_jpg': True
        },
        'Microsoft Excel Documents': {
            'aliases': ['Documentos do Microsoft Excel', 'Microsoft Excel Documents'],
            'title': 'Documentos do Microsoft Excel',
            'artifact_type': 'excel',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Nome do arquivo', ('Filename', 'Nome do arquivo')),
                ('Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Last Modified Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Created Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Tamanho (Bytes)', ('Size (Bytes)', 'File Size (Bytes)', 'Tamanho (Bytes)')),
                ('Autores', ('Authors', 'Autores')),
                ('Último autor', ('Last Author', 'Último autor')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': ['Arquivo', 'File'],
            'multi_seq_jpg': True
        },
        'CSV Documents': {
            'aliases': ['CSV Documents', 'Documentos CSV'],
            'title': 'Documentos CSV',
            'artifact_type': 'csv',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Nome do arquivo', ('File Name', 'Nome do arquivo')),
                ('Data/hora da modificação - UTC+00:00 (dd/MM/yyyy)',
                 ('Modified Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora da modificação - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora da criação - UTC+00:00 (dd/MM/yyyy)',
                 ('Created Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora da criação - UTC+00:00 (dd/MM/yyyy)')),
                ('Tamanho (Bytes)', ('Size (Bytes)', 'Tamanho (Bytes)')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            # CSV é tratado como texto: sem imagens geradas, apenas exibe o conteúdo em tabela
            'attachment_names': ['File Content', 'Conteúdo do arquivo'],
            'multi_seq_jpg': False
        },
        'Microsoft Sticky Notes': {
            'aliases': ['Microsoft Sticky Notes', 'Notas Autoadesivas do Microsoft'],
            'title': 'Notas Autoadesivas (Sticky Notes)',
            'artifact_type': 'sticky_notes',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Conteúdo', ('Content', 'Conteúdo')),
                ('Data/hora da criação - UTC+00:00 (dd/MM/yyyy)',
                 ('Created Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora da criação - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora da atualização - UTC+00:00 (dd/MM/yyyy)',
                 ('Updated Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora da atualização - UTC+00:00 (dd/MM/yyyy)')),
                ('ID do usuário', ('User ID', 'ID do usuário')),
                ('Estado', ('State', 'Estado')),
                ('Tema', ('Theme', 'Tema')),
                ('Posição', ('Position', 'Posição')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': [],
            'multi_seq_jpg': False
        },
        'Microsoft PowerPoint Documents': {
            'aliases': ['Microsoft PowerPoint Documents', 'Documentos do Microsoft PowerPoint'],
            'title': 'Documentos do Microsoft PowerPoint',
            'artifact_type': 'powerpoint',
            'fields': [
                ('Registro', ('sequenceNumber',)),
                ('Nome do arquivo', ('Filename', 'Nome do arquivo')),
                ('Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Last Modified Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora local da última modificação ao sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)',
                 ('File System Created Date/Time - UTC+00:00 (dd/MM/yyyy)',
                  'Data/hora de criação do sistema de arquivo - UTC+00:00 (dd/MM/yyyy)')),
                ('Tamanho (Bytes)', ('Size (Bytes)', 'File Size (Bytes)', 'Tamanho (Bytes)')),
                ('Autores', ('Authors', 'Autores')),
                ('Último autor', ('Last Author', 'Último autor')),
                ('Hash MD5', ('MD5 Hash', 'Hash MD5')),
                ('Hash SHA1', ('SHA1 Hash', 'Hash SHA1')),
                ('Fonte', ('Source', 'Fonte')),
            ],
            'attachment_names': ['File', 'Arquivo'],
            'multi_seq_jpg': False
        },
    }

# -------------------------
# Pipeline (um DOCX por artefato)
# -------------------------
def init_doc():
    doc = Document()
    sect = doc.sections[0]
    sect.page_height = Mm(297)
    sect.page_width  = Mm(210)
    sect.top_margin    = Mm(20)
    sect.bottom_margin = Mm(20)
    sect.left_margin   = Mm(25)
    sect.right_margin  = Mm(25)
    return doc

def add_title(doc, title_pt):
    p = doc.add_paragraph(title_pt)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def process_xml_to_multi_docx(xml_path, out_dir, log_callback=None, progress_callback=None, max_pages=0):
    mapping = build_mapping()
    xml_dir = os.path.dirname(xml_path)
    root = ET.parse(xml_path).getroot()
    artifacts = root.find('Artifacts')
    if artifacts is None:
        raise RuntimeError("XML sem <Artifacts>")

    groups = {}
    for artifact in artifacts.findall('Artifact'):
        name = artifact.get('name', '')
        found_map = None
        for _, info in mapping.items():
            if name in info['aliases']:
                found_map = info
                break
        if not found_map:
            continue
        hits = artifact.findall('.//Hit')
        if not hits:
            continue
        title = found_map['title']
        groups.setdefault(title, {'info': found_map, 'hits': []})
        groups[title]['hits'].extend(hits)

    total_hits = sum(len(g['hits']) for g in groups.values())
    done = 0
    out_files = []
    base = os.path.splitext(os.path.basename(xml_path))[0]

    for title, data in groups.items():
        info = data['info']
        hits = data['hits']

        doc = init_doc()
        add_title(doc, title)

        if log_callback:
            log_callback(f"Iniciando artefato: {title} ({len(hits)} registros).")

        for hit in hits:
            try:
                # ----- Tabela de metadados -----
                tbl = doc.add_table(rows=0, cols=2)
                tbl.autofit = False

                for label, aliases in info['fields']:
                    val = get_field_value(hit, aliases)
                    row = tbl.add_row().cells
                    row[0].text = str(label)
                    row[1].text = str(val if val is not None else '')

                loc = find_fragment_text(hit, 'Location', 'Localização')
                if loc and loc.strip().lower() != 'n/a':
                    row = tbl.add_row().cells
                    row[0].text = 'Localização'
                    row[1].text = loc

                rec = find_fragment_text(hit, 'Recovery method', 'Método de recuperação')
                if rec and rec.strip().lower() != 'n/a':
                    row = tbl.add_row().cells
                    row[0].text = 'Método de recuperação'
                    row[1].text = rec

                set_col_widths(tbl, [45, 115])
                style_table(tbl)

                spacer_tbl = doc.add_paragraph()
                spacer_tbl.paragraph_format.space_before = Pt(0)
                spacer_tbl.paragraph_format.space_after  = Pt(0)

                # ----- Anexos -----
                attachments_inserted = 0
                seq = hit.get('sequenceNumber', '')
                tmp_dir = os.path.join(out_dir, f"_tmp_imgs_{safe_slug(title)}_{seq}")

                artifact_type = info.get('artifact_type', '')

                if artifact_type == 'pdf':
                    # ── Documentos PDF ──
                    pdf_path = resolve_pdf_by_hyperlink(hit, xml_dir, log_callback)
                    if pdf_path:
                        imgs = generate_images_from_pdf(pdf_path, tmp_dir, dpi=200, max_pages=max_pages)
                        for fp in imgs:
                            insert_doc_image(doc, fp, max_w_mm=155, max_h_mm=215)
                            attachments_inserted += 1
                        if not imgs and log_callback:
                            log_callback(f"Aviso: PDF encontrado mas sem imagens geradas (Hit {seq}): {pdf_path}")
                    else:
                        if log_callback:
                            log_callback(f"Aviso: PDF não localizado para Hit {seq}.")

                elif artifact_type == 'word':
                    # ── Documentos Word ──
                    base_txt = find_fragment_text(hit, *info['attachment_names'])
                    if base_txt:
                        abs_docx = os.path.join(xml_dir, _normalize_relpath(base_txt))
                        imgs, warn = generate_images_from_word(abs_docx, tmp_dir, dpi=200, max_pages=max_pages)
                        if warn and log_callback:
                            log_callback(f"  {warn}")
                        for fp in imgs:
                            insert_doc_image(doc, fp, max_w_mm=155, max_h_mm=215)
                            attachments_inserted += 1
                    else:
                        if log_callback:
                            log_callback(f"Aviso: Fragment 'Arquivo/File' ausente no Hit {seq}.")

                elif artifact_type == 'excel':
                    # ── Documentos Excel ──
                    base_txt = find_fragment_text(hit, *info['attachment_names'])
                    if base_txt:
                        abs_xlsx = os.path.join(xml_dir, _normalize_relpath(base_txt))
                        imgs = generate_images_from_excel(abs_xlsx, tmp_dir, dpi=200,
                                                          max_pages=max_pages, log_callback=log_callback)
                        for fp in imgs:
                            insert_doc_image(doc, fp, max_w_mm=155, max_h_mm=215)
                            attachments_inserted += 1
                    else:
                        if log_callback:
                            log_callback(f"Aviso: Fragment 'Arquivo/File' ausente no Hit {seq}.")

                elif artifact_type == 'powerpoint':
                    # ── Documentos PowerPoint ──
                    base_txt = find_fragment_text(hit, *info['attachment_names'])
                    if base_txt:
                        abs_pptx = os.path.join(xml_dir, _normalize_relpath(base_txt))
                        imgs, warn = generate_images_from_powerpoint(
                            abs_pptx, tmp_dir, dpi=200, max_pages=max_pages, log_callback=log_callback
                        )
                        if warn and log_callback:
                            log_callback(f"  {warn}")
                        for fp in imgs:
                            insert_doc_image(doc, fp, max_w_mm=155, max_h_mm=215)
                            attachments_inserted += 1
                    else:
                        if log_callback:
                            log_callback(f"Aviso: Fragment 'File/Arquivo' ausente no Hit {seq}.")

                elif artifact_type == 'csv':
                    # ── Documentos CSV — insere conteúdo como tabela ──
                    # O caminho vem do Fragment attachment="True" (File Content)
                    csv_frag = None
                    for frag in hit.findall('Fragment'):
                        if frag.get('name') in ('File Content', 'Conteúdo do arquivo') or \
                           frag.get('attachment', '').lower() == 'true':
                            csv_frag = (frag.text or '').strip()
                            if csv_frag:
                                break
                    if csv_frag:
                        abs_csv = os.path.join(xml_dir, _normalize_relpath(csv_frag))
                        insert_csv_content(doc, abs_csv)
                        attachments_inserted += 1
                    else:
                        if log_callback:
                            log_callback(f"Aviso: caminho do CSV ausente no Hit {seq}.")

                elif artifact_type == 'sticky_notes':
                    # ── Sticky Notes — sem anexo de arquivo; conteúdo já está nos fields ──
                    pass  # Todos os dados relevantes já foram inseridos na tabela de metadados
                    # ── Fallback legado: sequência de JPEGs pré-gerados ──
                    base_txt = find_fragment_text(hit, *info['attachment_names'])
                    if base_txt:
                        abs_base = os.path.join(xml_dir, _normalize_relpath(base_txt))
                        imgs = list_sequential_jpegs(abs_base)
                        if not imgs and log_callback:
                            log_callback(f"Aviso: nenhum .jpg sequencial para base '{abs_base}' (Hit {seq}).")
                        for fp in imgs:
                            insert_doc_image(doc, fp, max_w_mm=155, max_h_mm=215)
                            attachments_inserted += 1
                    else:
                        if log_callback:
                            log_callback(f"Aviso: Fragment 'Arquivo/File' ausente no Hit {seq}.")

                else:
                    # ── Imagens/Thumbcache/Vídeos ──
                    frag_path = find_fragment_path(hit, info['attachment_names'])
                    if frag_path:
                        abs_path = os.path.join(xml_dir, _normalize_relpath(frag_path))
                        if os.path.isfile(abs_path):
                            insert_legacy_image(doc, abs_path, max_w_mm=155, max_h_mm=215)
                            attachments_inserted += 1
                        else:
                            if log_callback:
                                log_callback(f"Atenção: arquivo não encontrado: {abs_path}")

                # Limpeza de temporários
                try:
                    if os.path.isdir(tmp_dir):
                        import shutil
                        shutil.rmtree(tmp_dir, ignore_errors=True)
                except Exception:
                    pass

                if attachments_inserted > 0:
                    spacer_att = doc.add_paragraph()
                    spacer_att.paragraph_format.space_before = Pt(0)
                    spacer_att.paragraph_format.space_after  = Pt(0)

                if log_callback:
                    log_callback(f"{title} - Registro {seq} processado.")

            except Exception as e:
                if log_callback:
                    log_callback(f"ERRO no Hit {hit.get('sequenceNumber')}: {str(e)}")

            done += 1
            if progress_callback and total_hits:
                progress_callback(done / total_hits)

        slug = safe_slug(title)
        out_docx = os.path.join(out_dir, f"{base}_{slug}.docx")
        doc.save(out_docx)
        out_files.append(out_docx)
        if log_callback:
            log_callback(f"✅ Salvo: {out_docx}")

    if not out_files and log_callback:
        log_callback("Nenhum artefato correspondente encontrado no XML.")
    return out_files

# -------------------------
# GUI
# -------------------------
def select_and_run(log_widget, progress_bar, max_pages_var):
    xml_path = filedialog.askopenfilename(title="Selecione o XML", filetypes=[("XML files", "*.xml")])
    if not xml_path:
        return

    # Valida e lê o número de páginas
    try:
        mp = int(max_pages_var.get())
        if mp < 0:
            raise ValueError
    except ValueError:
        messagebox.showerror("Erro", "Número máximo de páginas deve ser um inteiro ≥ 0 (0 = todas).")
        return

    def set_progress(value):
        progress_bar['value'] = value * 100
        progress_bar.update()

    try:
        def logger(msg):
            log_widget.insert(tk.END, msg + "\n")
            log_widget.see(tk.END)
            log_widget.update()

        out_files = process_xml_to_multi_docx(
            xml_path,
            os.path.dirname(xml_path),
            log_callback=logger,
            progress_callback=set_progress,
            max_pages=mp
        )
        if out_files:
            log_widget.insert(tk.END, "\n✅ Concluído. Arquivos gerados:\n" + "\n".join(out_files) + "\n")
            progress_bar['value'] = 100
            if messagebox.askyesno("Abrir pasta?", f"Foram gerados {len(out_files)} arquivos DOCX.\n\nDeseja abrir a pasta agora?"):
                folder = os.path.dirname(out_files[0])
                try:
                    os.startfile(folder)
                except Exception:
                    import subprocess, sys
                    if sys.platform.startswith('darwin'):
                        subprocess.call(['open', folder])
                    else:
                        subprocess.call(['xdg-open', folder])
        else:
            messagebox.showinfo("Informação", "Nenhum artefato correspondente foi encontrado no XML.")
    except Exception as e:
        messagebox.showerror("Erro", str(e))
    finally:
        progress_bar['value'] = 0

def main_gui():
    root = tk.Tk()
    root.title("AXIOM XML → DOCX by Bogdan")
    root.geometry("850x660")

    frame = tk.Frame(root)
    frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

    lbl_info = tk.Label(
        frame,
        text=("📄 Instruções:\n"
              "1. Selecione um arquivo XML exportado pelo AXIOM (PT/EN).\n"
              "2. O programa criará um DOCX para CADA artefato (Imagens, Thumbcache, Vídeos, Word, PowerPoint, PDF, Excel, csv e Sticky Notes).\n"
              "3. PDFs são buscados automaticamente pelo caminho do XML; Word, PowerPoint e csv são convertidos se necessário.\n"
              "4. Devem ser gerados previamente os PDFs correspondentes aos arquivos de Excel e salvos no formato nome_do_arquivo.xlsx.pdf.\n"
              "5. Os arquivos são salvos no mesmo diretório do XML."),
        justify="left", anchor="w", font=("Segoe UI", 10)
    )
    lbl_info.pack(fill=tk.X, pady=(0, 8))

    # ── Controle de páginas ──
    pages_frame = tk.Frame(frame)
    pages_frame.pack(fill=tk.X, pady=(0, 8))

    tk.Label(pages_frame, text="Máximo de páginas por documento  (0 = todas):",
             font=("Segoe UI", 10)).pack(side=tk.LEFT)

    max_pages_var = tk.StringVar(value="0")
    pages_entry = tk.Entry(pages_frame, textvariable=max_pages_var, width=6,
                           font=("Segoe UI", 10), justify='center')
    pages_entry.pack(side=tk.LEFT, padx=(8, 0))

    tk.Label(pages_frame, text="páginas", font=("Segoe UI", 10)).pack(side=tk.LEFT, padx=(4, 0))

    # ── Botão principal ──
    btn = tk.Button(
        frame,
        text="Selecionar XML e gerar DOCX (um por artefato)",
        command=lambda: select_and_run(log, progress, max_pages_var)
    )
    btn.pack(fill=tk.X)

    progress = ttk.Progressbar(frame, orient='horizontal', mode='determinate')
    progress.pack(fill=tk.X, pady=5)

    log = scrolledtext.ScrolledText(frame, height=26)
    log.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

    root.mainloop()


if __name__ == "__main__":
    main_gui()